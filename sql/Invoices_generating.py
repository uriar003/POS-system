import docx
from docx import Document
from docx.shared import Inches
from docx import Document
from docx.shared import Pt, RGBColor
from SQL_Database import search_customer
from SQL_Database import see_item_bought
from SQL_Database import search_order
from SQL_Database import see_items
from SQL_Database import conn
#import win32com.client
import sqlite3
from datetime import datetime
from pathlib import Path
import sys, os
import json

if getattr(sys, 'frozen', False):
    # If the application is run as a bundle, the PyInstaller bootloader
    # extends the sys module by a flag frozen=True and sets the app 
    # path into variable _MEIPASS'.
    PARENTDIR = sys._MEIPASS
    with open(PARENTDIR+"/settings.json", "r") as fn:
        db = json.load(fn)
else:
    PARENTDIR = os.path.dirname(os.path.abspath(__file__))
    PARENTDIR = os.getcwd()
    i = PARENTDIR.rfind('/')
    PARENTDIR = PARENTDIR[:i] + "/json"
    PARENTDIR = str(Path(__file__).resolve().parent.parent)
    with open(PARENTDIR+"/json/settings.json", "r") as fn:
        db = json.load(fn)

#DATABASE_FILE = str(Path(__file__).resolve().parent)+"/POS_database.db"
# Opening of the database
'''
________________________________________________________________________________________________________________________
_______________________________________________READ ME__________________________________________________________________
You have to download 3 modules to create documents with python and interact with the Windows API:
->python-docx
->pywin32
->docx

If you don't have any EDU, let's see the following lines:
pip install pandas python-docx pywin32
pip install python-docx
________________________________________________________________________________________________________________________
________________________________________________________________________________________________________________________
'''


def make_client_invoice(name,email,order,list_item):
    list_number=[]
    for element in order:
        list_number.append(element[3])
    date=(order[0])[1]
    date=date.split(' ')
    date=date[0]
    date=date.split('-')
    date='/'.join(date)
    total_price=0
    #order=str(order)
    #chars = ['[', ']', "'", "(", ")"]
    #ord = order.translate(str.maketrans('', '', ''.join(chars)))
    #order=ord.split(',')
    document=Document()
    document.add_picture(PARENTDIR+'/brand_logo.png', width=Inches(1))
    document.add_paragraph('POS System Team')
    document.add_paragraph('333 S Twin Oaks Valley Rd')
    document.add_paragraph('San Marcos, CA 92096')
    document.add_heading('Receipt',0)
    p1= document.add_paragraph('Dear ')
    p1.add_run(name).bold=True
    p1.add_run(",")

    p2 = document.add_paragraph('Please find attached invoice for your recent purchase of the ')
    p2.add_run(str(date)).bold=True
    #p2.add_run(' units of ')
    #p2.add_run(product).bold=True
    #p2.add_run('.')

    p3 = document.add_paragraph('Invoice number: ')
    p3.add_run(str((order[0])[0])).bold=True

    [document.add_paragraph('') for _ in range(2)]

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item_id'
    hdr_cells[1].text = 'Product Name'
    hdr_cells[2].text = 'Units'
    hdr_cells[3].text = 'Unit Price ($)'
    hdr_cells[4].text = 'CA taxes 7.75% ($)'
    hdr_cells[5].text = 'Total Price ($)'
    list_item2=[]
    for element in list_item:
        chars = ['[', ']', "'", "(", ")"]
        res = str(element).translate(str.maketrans('', '', ''.join(chars)))
        list_item2.append(res)
    n=0
    for element in list_item2:
        element=element.split(',')
        print(element)
        for i in range(1):
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        row_cells = table.add_row().cells
        row_cells[0].text = str(element[0])
        row_cells[1].text = element[1]
        row_cells[2].text = f'{int(list_number[n]):,.2f}'
        row_cells[3].text = f'{float(element[3]):,.2f}'
        row_cells[4].text = f'{(int(list_number[n]) * float(element[3]))*(7.75/100):,.2f}'
        row_cells[5].text = f'{int(list_number[n])* float(element[3])+(int(list_number[n]) * float(element[3]))*(7.75/100):,.2f}'
        total_price=total_price + (int(list_number[n])* float(element[3])+(int(list_number[n]) * float(element[3]))*(7.75/100))
        total_price=round(total_price,2)
        [document.add_paragraph('') for _ in range(1)]
        n=n+1

    document.add_heading('Total Price', 0)
    p4 = document.add_paragraph('$ ')
    p4.add_run(str(total_price)).bold = True


    [document.add_paragraph('') for _ in range(1)]
    document.add_paragraph('We appreciate your business see you later!')
    document.add_paragraph('Sincerely')
    document.add_paragraph('POS system team')
    ourDoc = db["MainDirectory"]+'/Inventory/receipt.docx'
    document.save(ourDoc)
    return ourDoc


def informations(transaction_id):
    # Opening of the database
    #conn = sqlite3.connect(DATABASE_FILE)
    items_bought = see_item_bought(transaction_id)
    list_item_id=[]
    list_item=[]
    for element in items_bought:
        list_item_id.append(element[2])
    for element in list_item_id:
        item=see_items(element)
        list_item.append(item)

    return items_bought,list_item

"""
def docx_to_pdf(src, dst):
    word = win32com.client.Dispatch("Word.Application")
    wdFormatPDF = 17
    doc = word.Documents.Open(src)
    doc.SaveAs(dst, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

def send_email(name, email, attachment):
    outlook = win32com.client.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)
    mail.To = to_addr
    mail.Subject = 'Invoice from PythonInOffice'
    mail.Body = f'Dear {name}, Please find attached invoice'
    mail.Attachments.Add(attachment)
    mail.Send()
"""

#items_bought, list_item=informations('1')
#make_client_invoice('Julien Toulon', 'toulon.julien@gmail.com',items_bought,list_item)

#src=r'C:\Users\Julien\Documents\CS436\Clients' + r'\receipt.docx'
#dst=r'C:\Users\Julien\Documents\CS436\Clients' + r'\receipt.pdf'
#docx_to_pdf(src,dst)



