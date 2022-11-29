import docx
from docx import Document
from docx.shared import Inches
from docx import Document
from docx.shared import Pt, RGBColor
from SQL_database import search_customer
from SQL_database import see_item_boughts
from SQL_database import search_order
from SQL_database import see_items
import win32com.client
import sqlite3
from datetime import datetime

'''
________________________________________________________________________________________________________________________
_______________________________________________READ ME__________________________________________________________________
You have to download 3 modules to create documents with python and interact with the Windows API:
->python-docx
->pywin32
->docx

If you don't have any EDU, let's see the following lines:
pip install pandas python-docx pywin32
pip install python-docx
________________________________________________________________________________________________________________________
________________________________________________________________________________________________________________________
'''


def make_client_invoice(name,customer_id,email,order):
    total_price=0
    order=str(order)
    chars = ['[', ']', "'", "(", ")"]
    ord = order.translate(str.maketrans('', '', ''.join(chars)))
    order=ord.split(',')
    document=Document()
    document.add_picture('brand_logo.png', width=Inches(1))
    document.add_paragraph('POS System Team')
    document.add_paragraph('333 S Twin Oaks Valley Rd')
    document.add_paragraph('San Marcos, CA 92096')
    document.add_heading('Receipt',0)
    p1= document.add_paragraph('Dear ')
    p1.add_run(name).bold=True
    p1.add_run(' (customer id: ')
    p1.add_run(str(customer_id)).bold = True
    p1.add_run(")")
    p1.add_run(",")

    p2 = document.add_paragraph('Please find attached invoice for your recent purchase of the ')
    p2.add_run(str(date)).bold=True
    #p2.add_run(' units of ')
    #p2.add_run(product).bold=True
    #p2.add_run('.')

    p3 = document.add_paragraph('Invoice number: ')
    p3.add_run(str(order[0])).bold=True

    [document.add_paragraph('') for _ in range(2)]

    table = document.add_table(rows=1, cols=6)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Item_id'
    hdr_cells[1].text = 'Product Name'
    hdr_cells[2].text = 'Units'
    hdr_cells[3].text = 'Unit Price ($)'
    hdr_cells[4].text = 'CA taxes 7.75% ($)'
    hdr_cells[5].text = 'Total Price ($)'
    list_item2=[]
    for element in list_item:
        chars = ['[', ']', "'", "(", ")"]
        res = str(element).translate(str.maketrans('', '', ''.join(chars)))
        list_item2.append(res)
    print(list_item2)
    for element in list_item2:
        element=element.split(',')
        print(element)
        for i in range(1):
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        row_cells = table.add_row().cells
        row_cells[0].text = str(element[0])
        row_cells[1].text = element[1]
        row_cells[2].text = f'{int(element[2]):,.2f}'
        row_cells[3].text = f'{int(element[3]):,.2f}'
        row_cells[4].text = f'{(int(element[2]) * int(element[3]))*(7.75/100):,.2f}'
        row_cells[5].text = f'{int(element[2])* int(element[3])+(int(element[2]) * int(element[3]))*(7.75/100):,.2f}'
        total_price=total_price + (int(element[2])* int(element[3])+(int(element[2]) * int(element[3]))*(7.75/100))
        [document.add_paragraph('') for _ in range(1)]

    document.add_heading('Total Price', 0)
    p4 = document.add_paragraph('$ ')
    p4.add_run(str(total_price)).bold = True


    [document.add_paragraph('') for _ in range(1)]
    document.add_paragraph('We appreciate your business see you later!')
    document.add_paragraph('Sincerely')
    document.add_paragraph('POS system team')

    document.save(f'receipt.docx')


def informations(customer_id):
    # Opening of the database
    conn = sqlite3.connect('POS_database.db')
    rows=search_customer(customer_id)
    rows=list(rows[0])
    first_name=rows[1]
    last_name=rows[2]
    name=first_name + ' ' + last_name
    email=rows[3]
    customer_id=rows[0]
    date = str(datetime.now())
    date = '"', date, '"'
    date = ''.join(date)
    date=date.split(' ')
    date=list(date[0])
    date2=[]
    for element in date:
        if element!='"' and element!='-':
            date2.append(element)
        else:
            continue
    date2.insert(4, '/')
    date2.insert(7, '/')
    date2=''.join(date2)
    items_bought=see_item_boughts(customer_id,date2,date2)
    order=search_order(customer_id,date2,date2)
    list_item=[]
    for element in items_bought:
        item_id=element[2]
        item=see_items(item_id)
        list_item.append(item)

    return name,email,customer_id,date2,list_item,order

def docx_to_pdf(src, dst):
    word = win32com.client.Dispatch("Word.Application")
    wdFormatPDF = 17
    doc = word.Documents.Open(src)
    doc.SaveAs(dst, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

def send_email(name, email, attachment):
    outlook = win32com.client.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)
    mail.To = to_addr
    mail.Subject = 'Invoice from PythonInOffice'
    mail.Body = f'Dear {name}, Please find attached invoice'
    mail.Attachments.Add(attachment)
    mail.Send()

name,email,customer_id,date,list_item,order=informations('1')

make_client_invoice(name, customer_id,email,order)
src=r'C:\Users\Julien\Documents\CS436\Clients' + r'\receipt.docx'
dst=r'C:\Users\Julien\Documents\CS436\Clients' + r'\receipt.pdf'
docx_to_pdf(src,dst)



